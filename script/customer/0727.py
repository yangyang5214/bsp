import os
import docx
import pandas as pd
import numpy as np
import aspose.words as aw

np.float = np.float64


def read_file(fname: str):
    print('read file: ' + fname)

    if fname.endswith(".docx"):
        r = []
        doc = docx.Document(fname)
        for para in doc.paragraphs:
            r.append(para.text)
        return '\n'.join(r)
    elif fname.endswith(".rtf"):
        doc = aw.Document(fname)
        r = doc.to_string(aw.SaveFormat.TEXT)
        arrs = r.split('\n')
        return '\n'.join(arrs[2:len(arrs) - 4])
    else:
        with open(fname, 'r') as f:
            return '\n'.join(f.readlines())


def read_dir():
    result = []
    dir_name = "/Users/beer/Downloads/yfzf"
    for f in os.listdir(dir_name):
        if f.startswith('.'):
            continue
        name = f.split(".")[0]
        content = read_file(os.path.join(dir_name, f))
        result.append([name, content])

    # 1.txt
    with open("1.txt", 'r') as f:
        lines = f.readlines()
        pre_index = 0
        for index in range(len(lines)):
            line = lines[index]
            if line == '\n':
                result.append([lines[pre_index], ''.join(lines[pre_index+1:index])])
                pre_index = index + 1

    return result


def main():
    data = read_dir()
    df = pd.DataFrame(data, columns=['name', ''])
    with pd.ExcelWriter('/tmp/2.xlsx', engine='openpyxl') as writer:
        df.to_excel(writer, index=False)


if __name__ == '__main__':
    # read_dir()
    main()
